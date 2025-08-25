import os
import time
import hashlib
import pdfplumber
from docx import Document as DocxDocument
from pympler import asizeof


class Document:

    def __init__(self, path, load_content=True):
        start_time = time.time()

        self.path = path
        self.directory = os.path.dirname(path)
        self.name, ext = os.path.splitext(os.path.basename(path))
        self.extension = ext[1:].lower()
        self.size = os.path.getsize(path)

        self.tags = set()
        self.summary = None
        self.category = None

        with open(path, "rb") as f:
            self.hash_binary = self.hash_file(f, "md5")
            self.metadata = self.load_metadata(f)
            f.seek(0)
            self.content = self.load_content(f, load_content)

        self.hash_path = self.hash(self.path, "md5")
        self.hash_content = self.hash(self.content or "", "md5")
        self.hash_title = self.hash(self.name, "md5")

        self.load_time = (time.time() - start_time) * 1000 # miliseconds

    def load_content(self, file=None, load_content=True):
        """
        If called from constructor, use provided file object.
        If called independently, open self.path.
        """
        if not load_content:
            return None

        if file is None:
            file = open(self.path, "rb")
            close_after = True
        else:
            close_after = False

        try:
            content = None
            if self.extension == "docx":
                doc = DocxDocument(file)
                content = "\n".join(p.text for p in doc.paragraphs)

            elif self.extension == "pdf":
                with pdfplumber.open(file) as pdf:
                    content = "\n".join(
                        page.extract_text()
                        for page in pdf.pages
                        if page.extract_text()
                    )

            self.content = content

            return content
        finally:
            if close_after:
                file.close()


    def load_metadata(self, file):
        
        if self.extension == "docx":
            doc = DocxDocument(file)
            props = doc.core_properties
            return {
                attr: getattr(props, attr)
                for attr in dir(props)
                if not attr.startswith("_") and not callable(getattr(props, attr))
            }

        elif self.extension == "pdf":
            with pdfplumber.open(file) as pdf:
                return dict(pdf.metadata)

        return {}

    def clear_content(self):
        self.content = None

    def get_memory_usage(self):
        return asizeof.asizeof(self) # bytes

    def hash(self, content, function="md5"):
        h = hashlib.new(function)
        if isinstance(content, str):
            h.update(content.encode("utf-8"))
        elif isinstance(content, bytes):
            h.update(content)
        else:
            h.update(str(content).encode("utf-8"))
        return h.hexdigest()

    def hash_file(self, file, function="md5"):
        h = hashlib.new(function)
        for chunk in iter(lambda: file.read(4096), b""):
            h.update(chunk)
        return h.hexdigest()

    def set_tags(self, tags: set):
        self.tags.update(tags)

    def set_summary(self, summary: str):
        self.summary = summary

    def set_category(self, category: str):
        self.category = category

    def __eq__(self, doc):
        return isinstance(doc, Document) and self.hash_binary == doc.hash_binary

    def __hash__(self):
        return int(self.hash_path, 16)
    
    def __repr__(self):
        return (
            f"<Document name='{self.name}' "
            f"ext='{self.extension}' "
            f"size={self.size} bytes "
            f"tags={len(self.tags)} "
            f"summary={'yes' if self.summary else 'no'}>"
        )

    def __str__(self):
        parts = [f"Document: {self.name}.{self.extension} ({self.size} bytes)"]
        if self.summary:
            parts.append(f"Summary: {self.summary}")
        if self.tags:
            parts.append(f"Tags: {', '.join(sorted(self.tags))}")
        if self.category:
            parts.append(f"Category: {self.category}")
        return "\n".join(parts)

