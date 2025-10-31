from __future__ import annotations

import os
import hashlib
from typing import Dict, Any, Optional

import docx
from PyPDF2 import PdfReader
try:
    from langdetect import detect
except Exception:
    detect = None

from src.features.text_preproc import (
    filename_tokens,
    drop_repeating_lines,
)

# File text extraction & lightweight metadata
#
# Responsibilities:
# - Extract visible text from PDF/DOCX/TXT with pragmatic fallbacks.
# - Surface a small, stable metadata subset usable for indexing.
# - Compute a raw-file SHA256.
# - Defer any semantic normalization and content hashing to downstream stages.


# Byte-level identity
def compute_file_hash(file_path: str) -> str:
    """
    Compute a SHA256 over raw file bytes.
    """
    hasher = hashlib.sha256()
    with open(file_path, "rb") as f:
        while True:
            chunk = f.read(8192)
            if not chunk:
                break
            hasher.update(chunk)
    return hasher.hexdigest()


# PDF (.pdf)
def extract_text_pdf(file_path: str) -> str:
    """
    Extract visible text from a PDF, page-by-page.
    """
    try:
        reader = PdfReader(file_path)
    except Exception:
        # Unreadable/encrypted or invalid PDFs = empty text
        return ""

    chunks: list[str] = []
    for page in getattr(reader, "pages", []):
        try:
            pg = page.extract_text() or ""
        except Exception:
            pg = ""
        # Normalize line breaks up front for downstream stability
        pg = pg.replace("\r\n", "\n").replace("\r", "\n")
        chunks.append(pg)
    return "\n".join(chunks)


# WordprocessingML (.docx)
def extract_text_docx(file_path: str) -> str:
    """
    Extract visible paragraph text from a DOCX.
    """
    doc = docx.Document(file_path)
    # Preserve paragraph boundaries.
    paras = [para.text for para in doc.paragraphs]
    return "\n".join(paras)


# Plain text (.txt)
def extract_text_txt(file_path: str) -> str:
    """
    Read UTF-8 text.
    """
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


# Optional language detection
def _safe_detect_language(text: str) -> Optional[str]:
    """
    Attempt to detect language, returns None if unavailable/undecidable.
    """
    try:
        if detect is None or not text or not text.strip():
            return None
        return detect(text)
    except Exception:
        return None


# Metadata helpers
def _pdf_core_meta(file_path: str) -> Dict[str, Any]:
    """
    Extract a minimal set of PDF document metadata.
    """
    meta: Dict[str, Any] = {}
    try:
        reader = PdfReader(file_path)
        info = getattr(reader, "metadata", None) or {}
        # Metadata keys can be prefixed, normalize comparisons.
        def _get(key: str) -> Optional[str]:
            for k, v in info.items():
                if isinstance(k, str) and k.strip("/").lower() == key:
                    return str(v)
            return None

        meta["title"] = _get("title")
        meta["author"] = _get("author")
        meta["created"] = _get("creationdate") or _get("created")
    except Exception:
        # Metadata not essential
        pass
    return meta


def _docx_core_meta(file_path: str) -> Dict[str, Any]:
    """
    Extract a minimal set of DOCX core properties.
    """
    meta: Dict[str, Any] = {}
    try:
        doc = docx.Document(file_path)
        cp = doc.core_properties
        meta["title"] = cp.title or None
        meta["author"] = cp.author or None
        meta["created"] = cp.created.isoformat() if getattr(cp, "created", None) else None
    except Exception:
        pass
    return meta


# Main entry point
def extract_document(file_path: str) -> Dict[str, Any]:
    """
    Extract visible text and lightweight metadata from a file.

    Strategy
    - Dispatch by extension (.pdf/.docx/.txt), other formats raise ValueError.
    - For PDFs, apply `drop_repeating_lines` to reduce header/footer noise.
    - Compute a raw-file SHA256 and basic file level attributes.
    - Defer normalized content hashing to a downstream stage
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        raw_text = extract_text_pdf(file_path)
        # PDFs often include per-page headers/footers.
        raw_text = drop_repeating_lines(raw_text, min_count=3, max_len=120)
        core_meta = _pdf_core_meta(file_path)
        mime = "application/pdf"

    elif ext == ".docx":
        raw_text = extract_text_docx(file_path)
        core_meta = _docx_core_meta(file_path)
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    elif ext == ".txt":
        raw_text = extract_text_txt(file_path)
        core_meta = {}
        mime = "text/plain"

    else:
        raise ValueError(f"Unsupported file format: {ext}")

    # Best-effort language hint.
    lang = _safe_detect_language(raw_text)

    fname = os.path.basename(file_path)
    metadata: Dict[str, Any] = {
        "filename": fname,
        "filename_tokens": filename_tokens(fname),
        "filepath": file_path,
        "hash": compute_file_hash(file_path),
        "filesize": os.path.getsize(file_path),
        "language": lang,
        "mime_type": mime,
        "title": core_meta.get("title"),
        "author": core_meta.get("author"),
        "created": core_meta.get("created"),
    }

    return {"raw_text": raw_text, "metadata": metadata}
